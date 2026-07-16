"""
Document Parser Module

Extracts text from various document formats (PDF, TXT, DOCX).

Supported Formats:
    - PDF: Page-by-page extraction with page numbers
    - TXT: UTF-8 and common encodings
    - DOCX: Paragraph extraction with structure preservation
"""

import logging
from typing import List, Optional

import pypdf
from docx import Document as DocxDocument
from docx.table import Table

logger = logging.getLogger(__name__)


class ParseError(Exception):
    """Raised when document parsing fails."""
    pass


class DocumentParser:
    """
    Multi-format document parser.

    Handles PDF, TXT, and DOCX formats with format-specific optimizations.
    """

    SUPPORTED_FORMATS = {"pdf", "txt", "docx"}
    MAX_FILE_SIZE_MB = 50

    @staticmethod
    def _validate_file_path(file_path: str) -> None:
        """
        Validate file path before processing.

        Args:
            file_path: Path to file

        Raises:
            ParseError: If validation fails
        """
        if not file_path or not file_path.strip():
            raise ParseError("File path cannot be empty")

        if not isinstance(file_path, str):
            raise ParseError("Invalid file path type: %s", type(file_path).__name__)

    @staticmethod
    def parse_pdf(file_path: str) -> List[str]:
        """
        Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            List of page texts, one per page
            Returns empty list if PDF is empty/unreadable

        Raises:
            ParseError: If PDF is corrupted or unsupported
        """
        try:
            DocumentParser._validate_file_path(file_path)
            logger.debug("Parsing PDF: %s", file_path)
            texts: List[str] = []
            try:
                with open(file_path, "rb") as f:
                    reader = pypdf.PdfReader(f)

                    if not reader.pages:
                        logger.warning("PDF has no pages: %s", file_path)
                        return []

                    for page_num, page in enumerate(reader.pages, 1):
                        try:
                            text = page.extract_text()
                            if text and text.strip():
                                marked_text = f"[Page {page_num}]\n{text}"
                                texts.append(marked_text)
                            else:
                                logger.debug("Page %d has no text", page_num)

                        except Exception:
                            logger.warning("Failed to extract page %d", page_num)
                            continue

                if not texts:
                    logger.warning("No text extracted from PDF: %s", file_path)
                    return []

                logger.info("PDF parsing complete file: %s with %d pages extracted)", file_path, len(texts))

                return texts

            except pypdf.PdfReadError as e:
                raise ParseError("PDF is corrupted or unsupported") from e
            except FileNotFoundError:
                raise ParseError("PDF file not found: %s", file_path)

        except ParseError:
            raise
        except Exception as e:
            logger.exception("Unexpected error parsing PDF")
            raise ParseError("PDF parsing failed") from e

    @staticmethod
    def parse_txt(file_path: str) -> List[str]:
        """
        Extract text from TXT file.

        Process:
            1. Attempt UTF-8 decoding
            2. Fallback to latin-1 if UTF-8 fails
            3. Return entire content as single entry

        Args:
            file_path: Path to TXT file

        Returns:
            List with single element (entire file content)

        Raises:
            ParseError: If file not found or unreadable
        """
        try:
            DocumentParser._validate_file_path(file_path)
            logger.debug("Parsing TXT: %s", file_path)
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
                    encoding_used = "UTF-8"

            except UnicodeDecodeError:
                logger.debug("UTF-8 decode failed, trying latin-1: %s", file_path)

                with open(file_path, "r", encoding="latin-1") as f:
                    content = f.read()
                    encoding_used = "latin-1"

            if not content or not content.strip():
                logger.warning("TXT file is empty: %s", file_path)
                return []

            logger.info(
                "TXT parsing complete for file path: %s with %d chars with encoding)",
                file_path,
                len(content),
                encoding_used
            )

            return [content]

        except FileNotFoundError:
            raise ParseError("TXT file not found: %s", file_path)
        except Exception as e:
            logger.error("Unexpected error parsing TXT")
            raise ParseError("TXT parsing failed") from e

    @staticmethod
    def parse_docx(file_path: str) -> List[str]:
        """
        Extract text from DOCX file.

        Process:
            1. Load DOCX document
            2. Extract paragraphs
            3. Extract tables (convert to text)
            4. Preserve structure
            5. Return as single text block

        Args:
            file_path: Path to DOCX file

        Returns:
            List with single element (entire document content)

        Raises:
            ParseError: If file not found or corrupted
        """
        try:
            DocumentParser._validate_file_path(file_path)
            logger.debug("Parsing DOCX: %s", file_path)
            try:
                doc = DocxDocument(file_path)
            except Exception as e:
                raise ParseError("Failed to open DOCX file") from e

            texts: List[str] = []

            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    texts.append(para.text)

            for table in doc.tables:
                try:
                    table_text = DocumentParser._extract_table_text(table)
                    if table_text:
                        texts.append(table_text)
                except Exception:
                    logger.warning("Failed to extract table")
                    continue

            if not texts:
                logger.warning("No text extracted from DOCX: %s", file_path)
                return []

            content = "\n\n".join(texts)

            logger.info("DOCX parsing complete for file_path: %s with %d chars)", file_path, len(content))

            return [content]

        except FileNotFoundError:
            raise ParseError("DOCX file not found: %s", file_path)
        except Exception as e:
            logger.error("Unexpected error parsing DOCX")
            raise ParseError("DOCX parsing failed") from e

    @staticmethod
    def _extract_table_text(table: Table) -> Optional[str]:
        """
        Extract text from DOCX table.

        Converts table to text format with rows/columns preserved.

        Args:
            table: python-docx Table object

        Returns:
            Text representation of table, or None if empty
        """
        rows = []

        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join(
                    p.text for p in cell.paragraphs if p.text
                )
                cells.append(cell_text)

            if any(cells):
                rows.append(" | ".join(cells))

        if rows:
            return "\n".join(rows)
        return None

    def parse(
            self,
            file_path: str,
            file_type: str
    ) -> List[str]:
        """
        Parse document in any supported format.

        Main entry point - dispatches to format-specific parser.

        Args:
            file_path: Path to document file
            file_type: File type (pdf, txt, docx)

        Returns:
            List of text chunks (format-specific):
            - PDF: List of page texts
            - TXT: Single-element list with full content
            - DOCX: Single-element list with full content

        Raises:
            ParseError: If file not found, format unsupported, or parsing fails
            ValueError: If inputs invalid
        """
        try:
            if not file_path or not file_type:
                raise ValueError("file_path and file_type required")

            file_type = file_type.lower().strip()

            if file_type not in self.SUPPORTED_FORMATS:
                raise ValueError("Unsupported file type: %s. Supported: %s", file_type, self.SUPPORTED_FORMATS)

            logger.info("Parsing %s: %s", file_type.upper(), file_path)

            if file_type == "pdf":
                return self.parse_pdf(file_path)
            elif file_type == "txt":
                return self.parse_txt(file_path)
            elif file_type == "docx":
                return self.parse_docx(file_path)
            else:
                raise ParseError("Unsupported file type: %s", file_type)

        except (ValueError, ParseError):
            raise
        except Exception as e:
            logger.error("Unexpected error in parse()")
            raise ParseError("Document parsing failed") from e
